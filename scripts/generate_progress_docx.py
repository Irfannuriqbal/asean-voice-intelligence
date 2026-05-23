from __future__ import annotations

import json
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable


@dataclass(frozen=True)
class ReportInputs:
    group_name: str
    lecturer_name: str
    meeting_info: str
    meeting_date: str
    progress_text: str
    issues_text: str
    next_steps_text: str
    members_raw: str


def _split_bullets(text: str) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []

    lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        line = re.sub(r"^[-•\u2022\*]\s+", "", line)
        if line:
            lines.append(line)

    if len(lines) > 1:
        return lines

    # If user wrote one paragraph, try splitting by separators.
    single = lines[0] if lines else text
    parts = [p.strip() for p in re.split(r"[;•\u2022]\s*", single) if p.strip()]
    if len(parts) > 1:
        return parts

    return [single.strip()]


def _parse_members(members_raw: str) -> list[str]:
    raw = (members_raw or "").strip()
    if not raw:
        return []

    # Common pattern in the user's input: "NIM - Name" repeated.
    # Split when a new NIM-like number starts.
    chunks = re.split(r"\s+(?=\d{6,})", raw)
    cleaned: list[str] = []
    for c in chunks:
        c = c.strip(" -\n\t")
        if c:
            cleaned.append(c)
    return cleaned


def _load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def build_docx(inputs: ReportInputs, project_root: Path) -> Path:
    from docx import Document  # type: ignore
    from docx.shared import Inches

    outputs_dir = project_root / "outputs"
    outputs_dir.mkdir(parents=True, exist_ok=True)

    train_report = _load_json(outputs_dir / "train_report.json")
    cm_path = outputs_dir / "confusion_matrix.png"

    def metric(name: str, digits: int = 4) -> str:
        value = train_report.get(name)
        if isinstance(value, (int, float)):
            return f"{value:.{digits}f}"
        return "-"

    filename = (
        f"Laporan_Progres_Asistensi_{inputs.group_name.replace(' ', '_')}_"
        f"{date.today().isoformat()}.docx"
    )
    out_path = outputs_dir / filename

    doc = Document()

    doc.add_heading("LAPORAN PROGRES ASISTENSI", level=0)
    doc.add_paragraph("Mata Kuliah: Pengenalan Ucapan dan Teks ke Ucapan (PTU)")
    doc.add_paragraph("Proyek: ASR (MFCC + MLP) & TTS Bahasa Indonesia — Keyword Negara ASEAN")

    info = doc.add_table(rows=0, cols=2)
    info.style = "Table Grid"

    def add_row(label: str, value: str) -> None:
        row = info.add_row().cells
        row[0].text = label
        row[1].text = value

    add_row("Kelompok", inputs.group_name)
    add_row("Dosen", inputs.lecturer_name)
    add_row("Pertemuan/Asistensi", inputs.meeting_info)
    add_row("Tanggal", inputs.meeting_date)

    doc.add_paragraph("")
    doc.add_heading("Anggota", level=1)
    members = _parse_members(inputs.members_raw)
    if members:
        for m in members:
            doc.add_paragraph(m, style="List Bullet")
    else:
        doc.add_paragraph("-")

    doc.add_heading("Progres", level=1)
    for item in _split_bullets(inputs.progress_text) or ["-"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Kendala / Catatan", level=1)
    for item in _split_bullets(inputs.issues_text) or ["-"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Ringkasan Evaluasi Model (Terakhir)", level=1)
    eval_table = doc.add_table(rows=0, cols=2)
    eval_table.style = "Table Grid"

    def add_metric_row(label: str, key: str, digits: int = 4) -> None:
        r = eval_table.add_row().cells
        r[0].text = label
        r[1].text = metric(key, digits=digits)

    add_metric_row("Accuracy", "accuracy")
    add_metric_row("Precision (macro)", "precision_macro")
    add_metric_row("Recall (macro)", "recall_macro")
    add_metric_row("F1 (macro)", "f1_macro")
    add_metric_row("Jumlah sample", "n_samples", digits=0)
    add_metric_row("Train / Test", "n_train", digits=0)

    # Fix last row for train/test combined
    try:
        n_train = int(train_report.get("n_train", 0))
        n_test = int(train_report.get("n_test", 0))
        eval_table.rows[-1].cells[1].text = f"{n_train} / {n_test}" if (n_train and n_test) else "-"
    except Exception:
        pass

    if cm_path.exists():
        doc.add_paragraph("")
        doc.add_paragraph("Lampiran: Confusion Matrix")
        doc.add_picture(str(cm_path), width=Inches(6.0))

    doc.add_heading("Rencana Next Step", level=1)
    for item in _split_bullets(inputs.next_steps_text) or ["-"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Feedback Dosen", level=1)
    doc.add_paragraph("(Diisi saat asistensi)")
    doc.add_paragraph("")
    doc.add_paragraph("- ")
    doc.add_paragraph("- ")
    doc.add_paragraph("- ")

    doc.add_paragraph("")
    sig = doc.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    sig.cell(0, 0).text = "Mengetahui,\nDosen"
    sig.cell(0, 1).text = "Yang melaporkan,\nKelompok"
    sig.cell(1, 0).text = "\n\n(____________________)"
    sig.cell(1, 1).text = f"\n\n( {inputs.group_name} )"

    doc.save(str(out_path))
    return out_path


def main() -> int:
    project_root = Path(__file__).resolve().parents[1]

    # Defaults based on current chat inputs. Adjust via editing this file if needed.
    inputs = ReportInputs(
        group_name="Kelompok D9",
        lecturer_name="Anisa Putri",
        meeting_info="Ke-1",
        meeting_date="20 Mei 2026",
        progress_text="menyesuaikan projek keseleruhaan aja",
        issues_text="dari dataset",
        next_steps_text="menambahkan asset",
        members_raw=(
            "152023141 - Parisan Apro "
            "152023157 - Reynal Toeloes Ritonga "
            "152023 - Muhammad Yusuf Islam "
            "152023185 - Richa Romadhoni "
            "152023188 - Ramdhani Angriawan P"
        ),
    )

    out_path = build_docx(inputs, project_root=project_root)
    print(f"OK: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
